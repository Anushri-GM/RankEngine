import io
import json
import docx
import pytest
from app.understanding.job_understanding import JobUnderstandingEngine
from app.understanding.candidate_understanding import CandidateUnderstandingEngine
from app.understanding.behavior_understanding import BehaviorUnderstanding
from app.understanding.validation_engine import ValidationEngine
from app.understanding.normalization_engine import NormalizationEngine
from app.understanding.skill_taxonomy import taxonomy
from app.understanding.loader import Loader
from app.understanding.parser_metadata import ParserMetadataEngine
from app.decision_engine.decision_engine import DecisionEngine
from app.core.config import settings

def create_mock_docx() -> bytes:
    """
    Creates a mock docx binary stream for job description testing.
    """
    doc = docx.Document()
    doc.add_paragraph("Title: Senior Backend Developer")
    doc.add_paragraph("Location: San Francisco, CA (Hybrid)")
    doc.add_paragraph("Industry: Finance / Fintech")
    doc.add_paragraph("Required Skills: Python, Docker, Kubernetes, PostgreSQL, Javascript")
    doc.add_paragraph("Nice to have: AWS, Redis")
    doc.add_paragraph("Responsibilities:")
    doc.add_paragraph("• Design scalable backend services")
    doc.add_paragraph("• Coordinate Continuous Integration pipelines")
    doc.add_paragraph("• Mentorship of junior engineering developers")
    doc.add_paragraph("Experience: 5+ years of experience")
    doc.add_paragraph("Education: Master of Science in Computer Engineering")
    
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

def test_job_understanding():
    engine = JobUnderstandingEngine()
    docx_bytes = create_mock_docx()
    text = engine.read_docx(docx_bytes)
    
    assert "Senior Backend Developer" in text
    
    job_intel = engine.extract_intelligence(text)
    
    assert job_intel.role.title == "Senior Backend Developer"
    assert job_intel.role.category == "Backend Engineer"
    assert job_intel.role.seniority == "Senior"
    assert "Python" in job_intel.required_skills
    assert "Kubernetes" in job_intel.required_skills
    # Check normalization/aliases matching (e.g. Postgresql was matched to PostgreSQL)
    assert "PostgreSQL" in job_intel.required_skills
    assert "AWS" in job_intel.preferred_skills
    assert job_intel.experience_requirement.years == 5
    assert job_intel.education_requirement.degree == "Masters"
    assert "Engineering" in job_intel.education_requirement.major
    # Inferred expectations:
    assert "Backend Architecture" in job_intel.inferred_expectations
    assert "REST APIs" in job_intel.inferred_expectations

def test_candidate_understanding_and_behavior():
    raw_candidates = Loader.get_sample_candidates()
    candidate = raw_candidates[0]
    
    cand_intel = CandidateUnderstandingEngine.process_candidate(candidate)
    
    assert cand_intel.name == "Sarah Jenkins"
    # Normalization of skills (e.g. node -> Node.js, amazon web services -> AWS)
    assert "Node.js" in cand_intel.skills
    assert "AWS" in cand_intel.skills
    # Experience calculation: 2021-01 to 2023-05 is 28 months, 2023-06 to present (e.g. 2026) is > 36 months
    assert cand_intel.experience_years > 2.0
    assert len(cand_intel.projects) == 1
    assert cand_intel.projects[0].business_domain == "Finance"
    assert cand_intel.education[0].degree == "Bachelors"
    assert cand_intel.certifications[0].provider == "AWS"
    
    # Behavior Profile Generation
    profile = BehaviorUnderstanding.generate_profile(cand_intel, "Valid")
    
    assert profile.profile_completeness_score == 100
    assert profile.validation_status == "Valid"
    assert profile.behavior_score > 60
    assert "Cloud" in profile.skill_categories
    assert "Backend" in profile.skill_categories

def test_validation_engine():
    # Valid candidate test
    raw_candidates = Loader.get_sample_candidates()
    report = ValidationEngine.validate_candidate_dataset(raw_candidates)
    
    assert report.status in ["Valid", "Valid with Warnings"]
    assert len(report.failures) == 0
    
    # Invalid candidate test (duplicate IDs and invalid dates)
    invalid_candidates = [
        {
            "candidate_id": "cand-01",
            "name": "Sarah Jenkins",
            "skills": ["React"],
            "work_history": [
                {
                    "company": "Tech Corp",
                    "role": "Software Developer",
                    "start_date": "2023-01",
                    "end_date": "2022-01" # Start after end
                }
            ]
        },
        {
            "candidate_id": "cand-01", # Duplicate ID
            "name": "Sarah Two",
            "skills": ["Python"]
        }
    ]
    
    report_invalid = ValidationEngine.validate_candidate_dataset(invalid_candidates)
    assert report_invalid.status == "Invalid"
    assert "cand-01" in report_invalid.duplicate_ids
    assert len(report_invalid.failures) > 0

def test_normalization_engine():
    assert NormalizationEngine.normalize_skill("Amazon Web Services") == "AWS"
    assert NormalizationEngine.normalize_skill("Javascript") == "JavaScript"
    assert NormalizationEngine.normalize_skill("ML") == "Machine Learning"
    assert NormalizationEngine.normalize_skill("Node") == "Node.js"
    assert NormalizationEngine.normalize_skills_list(["node", "Javascript"]) == ["Node.js", "JavaScript"]

def test_taxonomy():
    assert taxonomy.get_canonical("k8s") == "Kubernetes"
    assert taxonomy.get_category("PostgreSQL") == "Database"
    assert taxonomy.get_category("React") == "Frontend"


def test_parser_metadata_generation(tmp_path):
    candidates = Loader.get_sample_candidates()
    processed_candidates = [CandidateUnderstandingEngine.process_candidate(c) for c in candidates]
    validation_report = ValidationEngine.validate_candidate_dataset(candidates)
    behavior_profiles = [BehaviorUnderstanding.generate_profile(candidate, validation_report.status) for candidate in processed_candidates]

    manifest, parser_report, skill_taxonomy, role_taxonomy, dataset_statistics = ParserMetadataEngine.generate_artifacts(
        candidates=candidates,
        processed_candidates=processed_candidates,
        behavior_profiles=behavior_profiles,
        validation_report=validation_report,
        dataset_path=tmp_path / "sample_candidates.json",
        output_dir=tmp_path / "outputs",
    )

    assert manifest["candidate_count"] == len(candidates)
    assert manifest["dataset_name"] == "sample_candidates"
    assert manifest["valid_candidates"] >= 1
    assert parser_report["average_skills"] > 0
    assert parser_report["average_experience"] > 0
    assert parser_report["average_projects"] > 0
    assert skill_taxonomy[0]["canonical_name"] in {"Python", "JavaScript", "React"}
    assert role_taxonomy[0]["role_category"]
    assert dataset_statistics["candidate_count"] == len(candidates)


def test_discover_candidate_dataset_from_parent_workspace(tmp_path):
    project_root = tmp_path / "reci-ai"
    project_root.mkdir()

    dataset_path = tmp_path / "Dataset" / "[PUB] India_runs_data_and_ai_challenge" / "India_runs_data_and_ai_challenge" / "sample_candidates.json"
    dataset_path.parent.mkdir(parents=True, exist_ok=True)
    dataset_path.write_text(json.dumps([{"candidate_id": "cand-1", "name": "Test Candidate", "skills": ["Python"]}]), encoding="utf-8")

    discovered_path, discovered_paths = ParserMetadataEngine.discover_candidate_dataset(base_dir=project_root)

    assert discovered_path == dataset_path.resolve()
    assert dataset_path.resolve() in discovered_paths


def test_pre_run_validate_outputs_generates_candidate_intelligence(tmp_path, monkeypatch):
    outputs_dir = tmp_path / "outputs"
    cache_dir = tmp_path / "cache"
    outputs_dir.mkdir(parents=True, exist_ok=True)
    cache_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(settings, "OUTPUT_PATH", str(outputs_dir))
    monkeypatch.setattr(settings, "MODEL_CACHE", str(cache_dir))

    dataset_path = tmp_path / "Dataset" / "sample_candidates.json"
    dataset_path.parent.mkdir(parents=True, exist_ok=True)
    dataset_path.write_text(json.dumps([{"candidate_id": "cand-1", "name": "Test Candidate", "skills": ["Python"]}]), encoding="utf-8")

    engine = DecisionEngine()
    engine.pre_run_validate_outputs()

    candidate_path = outputs_dir / "candidate_intelligence.json"
    assert candidate_path.exists()

    payload = json.loads(candidate_path.read_text(encoding="utf-8"))
    assert len(payload) == 1
    assert payload[0]["candidate_id"] == "cand-1"
