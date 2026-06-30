import io
import re
import docx
import spacy
from typing import List
from app.understanding.schema import JobIntelligence, Role, ExperienceRequirement, EducationRequirement
from app.understanding.role_understanding import RoleUnderstanding
from app.understanding.skill_taxonomy import taxonomy

class JobUnderstandingEngine:
    def __init__(self):
        # Attempt to load spaCy model, fall back gracefully if missing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception:
            self.nlp = None

    def read_docx(self, docx_bytes: bytes) -> str:
        """
        Extracts all paragraph and table text from a raw docx byte stream.
        """
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)

    def extract_intelligence(self, text: str) -> JobIntelligence:
        """
        Processes text and extracts standard metadata, required skills, and infers expectations.
        """
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # 1. Title Extraction
        title = "Software Engineer"
        for line in lines[:5]:
            match = re.search(r'(?:title|role|position)\s*:\s*(.*)', line, re.IGNORECASE)
            if match:
                title = match.group(1).strip()
                break
        else:
            if lines:
                title = lines[0]

        # 2. Skill Extraction (Canonicalization via Taxonomy)
        all_skills = []
        for s in taxonomy.skills:
            for alias in s.aliases:
                if re.search(r'\b' + re.escape(alias) + r'\b', text, re.IGNORECASE):
                    all_skills.append(s.canonical_name)
                    break
        
        # Distribute into Required and Preferred based on sections
        text_lower = text.lower()
        pref_section = ""
        
        pref_headers = ["preferred", "nice to have", "plus", "desirable", "bonus", "beneficial"]
        for header in pref_headers:
            idx = text_lower.find(header)
            if idx != -1:
                pref_section = text[idx:]
                break
                
        required_skills = []
        preferred_skills = []
        
        for skill in set(all_skills):
            if pref_section and re.search(r'\b' + re.escape(skill) + r'\b', pref_section, re.IGNORECASE):
                preferred_skills.append(skill)
            else:
                required_skills.append(skill)

        # 3. Responsibilities
        responsibilities = []
        in_resp = False
        for line in lines:
            line_lower = line.lower()
            if any(w in line_lower for w in ["responsibilities", "what you will do", "role description", "key tasks", "duty", "duties"]):
                in_resp = True
                continue
            if in_resp and any(w in line_lower for w in ["requirements", "skills", "qualifications", "what we offer", "benefits"]):
                in_resp = False
            
            if in_resp and (line.startswith("•") or line.startswith("-") or line.startswith("*") or re.match(r'^\d+\.', line)):
                responsibilities.append(re.sub(r'^[•\-\*\d\.\s]+', '', line).strip())

        if not responsibilities:
            # Bullet extractors fallback
            for line in lines:
                if line.startswith("•") or line.startswith("-") or line.startswith("*"):
                    responsibilities.append(re.sub(r'^[•\-\*\s]+', '', line).strip())
            responsibilities = responsibilities[:6]

        # 4. Experience Requirement
        exp_years = 2
        exp_desc = "2+ years of experience required"
        exp_match = re.search(r'(\d+)\s*\+?\s*years?(?:\s*of)?\s*(?:relevant\s*)?experience', text, re.IGNORECASE)
        if exp_match:
            exp_years = int(exp_match.group(1))
            exp_desc = f"{exp_years}+ years of experience required"

        # 5. Education Requirement
        degree = "Bachelors"
        if "phd" in text_lower or "ph.d" in text_lower or "doctorate" in text_lower:
            degree = "Doctorate"
        elif "master" in text_lower or "m.s." in text_lower or "msc" in text_lower or "mba" in text_lower:
            degree = "Masters"
        
        majors = []
        if "computer science" in text_lower or "cs" in text_lower:
            majors.append("Computer Science")
        if "engineering" in text_lower:
            majors.append("Engineering")
        if "mathematics" in text_lower or "stats" in text_lower or "statistics" in text_lower:
            majors.append("Mathematics / Statistics")
        if not majors:
            majors.append("Computer Science")

        # 6. Classification & Metadata
        role_category, seniority = RoleUnderstanding.classify_role(title, required_skills)
        role_obj = Role(title=title, category=role_category, seniority=seniority)

        industry = "Technology"
        if any(w in text_lower for w in ["finance", "fintech", "bank", "trading", "capital"]):
            industry = "Finance"
        elif any(w in text_lower for w in ["medical", "health", "hospital", "pharma", "clinical"]):
            industry = "Healthcare"
        elif any(w in text_lower for w in ["ecommerce", "e-commerce", "retail", "shop", "logistics"]):
            industry = "E-commerce"

        employment_type = "Full-time"
        if "contract" in text_lower:
            employment_type = "Contract"
        elif "part-time" in text_lower or "part time" in text_lower:
            employment_type = "Part-time"
        elif "internship" in text_lower or "intern" in text_lower:
            employment_type = "Internship"

        location = "Remote"
        loc_match = re.search(r'(?:location|based in|office)\s*:\s*([^\n]+)', text, re.IGNORECASE)
        if loc_match:
            location = loc_match.group(1).strip()
        elif "hybrid" in text_lower:
            location = "Hybrid"
        elif "on-site" in text_lower or "onsite" in text_lower:
            location = "On-site"

        # 7. Soft Skills
        soft_skills_vocab = ["communication", "teamwork", "leadership", "problem solving", "critical thinking", "collaboration", "adaptability", "mentorship", "creativity"]
        soft_skills = [skill.capitalize() for skill in soft_skills_vocab if skill in text_lower]

        # 8. Keywords & POS analysis
        keywords = []
        if self.nlp:
            doc = self.nlp(text)
            nouns = [chunk.text for chunk in doc.noun_chunks if len(chunk.text.split()) < 3 and len(chunk.text) > 3]
            keywords = list(set(nouns))[:8]
        else:
            keywords = [w.capitalize() for w in title.split()] + required_skills[:4]

        # 9. Inferred Hidden Expectations (Step 2)
        inferred = []
        if any(w in text_lower for w in ["scalable", "scalability", "distributed", "backend", "microservices", "concurrency"]):
            inferred.extend(["Backend Architecture", "REST APIs", "Distributed Systems", "Microservices"])
        if any(w in text_lower for w in ["machine learning", "ml", "models", "nlp", "ai", "inference"]):
            inferred.extend(["Model Evaluation", "ML Pipeline Design", "Feature Engineering", "Neural Networks"])
        if any(w in text_lower for w in ["react", "frontend", "ui", "ux", "browser", "dom"]):
            inferred.extend(["State Management", "DOM Performance", "Web Accessibility (a11y)", "Responsive Grid Systems"])
        if any(w in text_lower for w in ["devops", "cloud", "aws", "kubernetes", "ci/cd", "infrastructure"]):
            inferred.extend(["Infrastructure as Code", "Container Orchestration", "Deployment Pipelines", "Zero Downtime Deploys"])
            
        inferred = list(set(inferred))

        return JobIntelligence(
            role=role_obj,
            industry=industry,
            required_skills=required_skills,
            preferred_skills=preferred_skills,
            responsibilities=responsibilities,
            experience_requirement=ExperienceRequirement(years=exp_years, description=exp_desc),
            education_requirement=EducationRequirement(degree=degree, major=majors),
            employment_type=employment_type,
            location=location,
            soft_skills=soft_skills,
            keywords=keywords,
            inferred_expectations=inferred
        )
